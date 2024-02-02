from docxtpl import InlineImage
from docx.shared import Inches
import os
from datetime import datetime
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml import OxmlElement

def save_image(image, path):
    if image:
        image.save(path)

def generate_image_paths(num_images):
    return [f'static/imagen{i}.png' for i in range(1, num_images + 1)]

def generate_inline_images(document, image_paths):
    return [InlineImage(document, path, width=Inches(7.5), height=Inches(4.0)) if os.path.exists(path) else None for path in image_paths]

def generate_date():
    date_time = datetime.now()
    return date_time.strftime("%d-%m-%Y")

def generate_time():
    date_time = datetime.now()
    return date_time.strftime("%H:%M:%S")

def configuration_word(doc):
    section = doc.sections[0]
    section.left_margin = Cm(1.27) 
    section.right_margin = Cm(1.27)  
    section.top_margin = Cm(1.27)  
    section.bottom_margin = Cm(1.27)
    
def draw_table(doc, responsable, aplicativo, producto, indicePrueba, descripcionPrueba, 
               observaciones, estado, fecha):
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'

    width_column_1 = Cm(4)
    width_column_2 = Cm(16)
        
    for i in range(8):
        cell_1 = table.cell(i, 0)
        cell_2 = table.cell(i, 1)

        cell_1.width = width_column_1
        cell_2.width = width_column_2

        cell_border = cell_1._element.xpath('.//w:tcBorders')
        if not cell_border:
            cell_border = [OxmlElement('w:tcBorders')]
            cell_1._element.append(cell_border[0])
        cell_border[0].append(OxmlElement('w:right'))

        headers = ["Responsable", "Aplicativo", "Producto", "Indice caso de prueba",
                   "Descripcion caso de prueba", "Observaciones", "Estado de la prueba", 
                   "Fecha de generación del reporte"]

        cell_header = table.cell(i, 0)
        cell_header.text = headers[i]
        cell_header.paragraphs[0].runs[0].bold = True

    data_document = [responsable, aplicativo, producto, indicePrueba, descripcionPrueba,
                        observaciones, estado, fecha]

    for i, value in enumerate(data_document):
        table.cell(i, 1).text = value

    table.alignment = 1
    doc.add_paragraph()
