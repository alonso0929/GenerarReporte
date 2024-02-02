from flask import Flask, render_template, request, send_file
from io import BytesIO
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Inches, Cm
from docx.shared import RGBColor
from bs4 import BeautifulSoup
from utils import generate_date, configuration_word, draw_table

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('indexV2.html')

@app.route('/generate_word', methods=['POST'])
def generate_word():
    responsable = request.form['responsable']
    aplicativo = request.form['aplicativo']
    producto = request.form['producto']
    indicePrueba = request.form['indiceprueba']
    descripcionPrueba = request.form['descripcionprueba']
    observaciones = request.form['observaciones']
    ejecucion = request.form['ejecucion']
    estado = request.form['estado']

    documento = Document()
    configuration_word(documento)

    documento.add_heading(f'REPORTE DE PRUEBAS FLAGON {aplicativo} - {producto}', 0)

    p_titulo = documento.add_paragraph()
    p_titulo.alignment = 1
    run_titulo = p_titulo.add_run('Hoja de Control')
    run_titulo.bold = True

    draw_table(documento, responsable, aplicativo, producto, indicePrueba, descripcionPrueba, observaciones,
               estado, generate_date())

    p_ejecucion = documento.add_paragraph()
    run_ejecucion = p_ejecucion.add_run(f'Evidencia ejecución de la prueba: {ejecucion}')
    run_ejecucion.bold = True

    soup = BeautifulSoup(render_template('indexV2.html'), 'html.parser')
    labels = soup.find_all('label', {'name': True})

    for i, label in enumerate(labels, start=1):
        texto = label.text.strip()
        opcion_estado = request.form.get(f'estado{i}', '')
        opcion_radio = request.form.get(f'opcion_radio{i}', '')
        opcion_comentarios = request.form.get(f'comentario{i}', '')

        p_op_radio = documento.add_paragraph()
        run_op_radio = p_op_radio.add_run(f'{texto} {opcion_radio}')
        run_op_radio.bold = True

        p_op_estado = documento.add_paragraph()
        run_op_estado = p_op_estado.add_run(f'{opcion_estado}')
        run_op_estado.bold = True

        if opcion_estado=="Conforme":
            run_op_radio.font.color.rgb = RGBColor(0, 255, 0)
            run_op_estado.font.color.rgb = RGBColor(0, 255, 0)
        elif opcion_estado=="Observado":
            run_op_radio.font.color.rgb = RGBColor(255, 0, 0)
            run_op_estado.font.color.rgb = RGBColor(255, 0, 0)
        else:
            print("No aplica")
        
        documento.add_paragraph(f'{opcion_comentarios}')

        input_name = f'imagenes{i}[]'
        imagenes = request.files.getlist(input_name)

        if imagenes:
            for imagen in imagenes:
                try:
                    documento.add_picture(imagen, width=Inches(7.5), height=Inches(4.0))
                except UnrecognizedImageError as e:
                    pass

    output_stream = BytesIO()
    documento.save(output_stream)
    output_stream.seek(0)

    return send_file(output_stream, as_attachment=True, download_name=f'EVIDENCIA QA {aplicativo} {producto} {indicePrueba}.docx')

if __name__ == '__main__':
    app.run(debug=True)
